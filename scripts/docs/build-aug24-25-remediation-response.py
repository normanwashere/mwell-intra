from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "outputs" / "Mwell_Intra_WMS_August_24_25_Remediation_Response.docx"
EVIDENCE = ROOT / "docs" / "evidence" / "2026-08-26-aug24-25-remediation"
LOGO = ROOT / "apps" / "shell" / "public" / "mwell-wordmark.png"

BLUE = "0069B4"
DEEP_BLUE = "0B2545"
CYAN = "31C1E6"
MUTED = "5D6B82"
LIGHT = "F2F6FA"
PALE_BLUE = "E8F3FA"
PALE_GREEN = "E8F7F1"
GREEN = "087A5B"
PALE_GOLD = "FFF6E4"
GOLD = "9A6100"
RED = "A61B33"
WHITE = "FFFFFF"
BLACK = "17233A"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_run(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold_lead=None, color=BLACK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run(rest, color=color)
    else:
        set_run(p.add_run(text), color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run(p.add_run(text))
    return p


def add_step(doc, number, title, detail):
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [0.55, 5.95])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    left, right = table.rows[0].cells
    set_cell_fill(left, BLUE)
    set_cell_fill(right, LIGHT)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(str(number)), size=12, bold=True, color=WHITE)
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title), bold=True, color=DEEP_BLUE)
    p.add_run("\n")
    set_run(p.add_run(detail), size=10, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_status_callout(doc, title, text, status="verified"):
    colors = {
        "verified": (PALE_GREEN, GREEN),
        "note": (PALE_BLUE, BLUE),
        "caution": (PALE_GOLD, GOLD),
    }
    fill, accent = colors[status]
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title.upper()), size=9, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), size=10.5, color=DEEP_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, DEEP_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=9, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_idx % 2:
                set_cell_fill(cells[idx], LIGHT)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run(p.add_run(str(value)), size=font_size, color=BLACK)
    return table


def add_figure(doc, filename, title, explanation, width=6.25):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(title), size=12, bold=True, color=DEEP_BLUE)
    image_path = EVIDENCE / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(explanation), size=9.5, italic=True, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DEEP_BLUE),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_document():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("mWell Intra  |  UAT Remediation Evidence"), size=8.5, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Internal UAT  |  August 26, 2026  |  Page "), size=8.5, color=MUTED)
    add_page_number(p)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.add_run().add_picture(str(LOGO), width=Inches(1.5))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run("REMEDIATION RESPONSE"), size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run("August 24–25 WMS Feedback"), size=25, bold=True, color=DEEP_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run(p.add_run("What changed, how the controlled workflows operate, and live UAT evidence"), size=13, color=MUTED)

    metadata = [
        ("Environment", "mwell-intra-uat.vercel.app"),
        ("Source", "wms comments (3).pdf — August 24 and August 25 feedback"),
        ("Prepared", "August 26, 2026"),
        ("Status", "Reported items implemented and re-verified in live UAT"),
    ]
    add_table(doc, ["Document control", "Value"], metadata, [1.35, 5.15], font_size=10)
    doc.add_paragraph()
    add_status_callout(
        doc,
        "Verified outcome",
        "The reported workflows are represented in live UAT, the deterministic scenario data remains available, and four role-authenticated evidence journeys pass against the deployed Vercel application.",
    )

    doc.add_heading("Executive result", level=1)
    add_body(doc, "The August 24–25 feedback was retested through the same roles and records described by the warehouse team. The implementation now provides a controlled path from Procurement PO to receiving and quality, from demand to pick/pack/release, and from Marketing request to event custody and reconciliation.")
    add_body(doc, "During evidence capture, one fixture attribution mismatch was found: Event A had been seeded under the General Employee actor. Least-privilege behavior correctly hid it from Marketing. Both the live record and reusable seeder now attribute it to the Marketing Events Lead, preserving privacy while giving that tester a visible starting record.")

    doc.add_heading("Scope and evidence method", level=1)
    for item in (
        "Reviewed the original stakeholder report, emphasizing comments dated August 24 and August 25.",
        "Authenticated as Operations Associate and Marketing Events Lead against the live Vercel UAT deployment.",
        "Captured the exact controls and records used by each workflow without submitting destructive transactions.",
        "Read the UAT Supabase project directly to confirm record counts and ownership.",
        "Ran deterministic fixture regression tests and visually inspected every screenshot included here.",
    ):
        add_bullet(doc, item)

    doc.add_heading("What was fixed", level=1)
    fixes = [
        ("Marketing request access", "Added New stock request, retained own-request visibility, and repaired Event A ownership to the Marketing actor."),
        ("Governed PO receiving", "Goods POs reach Operations with per-line clean, damaged, unidentified, short, and excess quantities plus evidence and staging context."),
        ("Quality and putaway", "Accepted stock follows inspection/putaway; non-clean quantities remain controlled exceptions instead of silently becoming available stock."),
        ("Fulfillment coverage", "Scenario queues span Ecommerce, Shopify, event sales, bundles, picking, packing, release, cancellation, delivery failure, and proof of delivery."),
        ("Returns and replacement", "Serial-linked intake, decision, replacement, resolution, customer evidence, and closure states are seeded."),
        ("Responsive usability", "Long task identifiers wrap cleanly at 320 px and role-specific actions retain legible hierarchy on desktop."),
    ]
    for index, (title, detail) in enumerate(fixes, 1):
        add_step(doc, index, title, detail)

    doc.add_page_break()
    doc.add_heading("Live UAT seed readiness", level=1)
    add_body(doc, "The following counts were read directly from the UAT Supabase project on August 26, 2026. They are intended to keep cross-role testing productive without requiring users to create every prerequisite manually.")
    seed_rows = [
        ("Procurement requests", 3, "Request-to-PO traceability"),
        ("Purchase orders", 3, "Two open receipt cases; one completed reference"),
        ("Purchase-order lines", 6, "Four-line device PO plus two supply lines"),
        ("Active warehouse bins", 7, "Device, supply, and quarantine putaway"),
        ("Scenario products", 10, "Devices, OTG supplies, paper bags, and lanyard"),
        ("Serialized inventory units", 140, "Serial lookup, allocation, picking, and returns"),
        ("Fulfillment orders", 12, "Active and terminal fulfillment stages"),
        ("Customer return cases", 4, "Intake through customer closure"),
        ("Event records / reconciliations", "1 / 1", "Event A custody and settlement"),
        ("Quality receipts / inspections", "1 / 1", "Completed PO reference"),
        ("Kit definitions", 1, "OTG bundle completeness"),
    ]
    add_table(doc, ["Scenario data", "Count", "Purpose"], seed_rows, [2.2, 0.65, 3.65], font_size=9.2)
    add_status_callout(doc, "Tester-safe starting points", "PO 0001 and PO 0002 remain open. PO 0003 is a completed reference. The guarded seeder restores missing deterministic rows without overwriting unrelated tester-created records.", "note")

    doc.add_heading("Feedback response matrix", level=1)
    matrix = [
        ("Marketing cannot request stock", "New stock request and governed form", "Fig. 3"),
        ("Marketing sees unrelated requests", "Own-request scope retained; Event A owner corrected", "Fig. 3 + DB"),
        ("PO 0001 mixed outcomes", "Five condition quantities per line", "Fig. 2"),
        ("PO 0002 supply receiving", "Open issued PO in Operations queue", "Fig. 1"),
        ("PO 0003 putaway feedback", "Completed receipt/inspection reference retained", "DB"),
        ("Rack/bin mapping", "Seven seeded bins including quarantine", "DB"),
        ("Ecommerce / Shopify", "Channel, allocation, scan, pack, release states", "Fig. 5"),
        ("Bundle sets", "OTG kit and set codes seeded", "DB"),
        ("Event A sale", "Event, third-party order, and reconciliation", "Figs. 4–5"),
        ("Returns / replacement", "Four lifecycle states with serial and evidence", "DB"),
        ("Split backorder", "Explicit order action", "Fig. 5"),
        ("Mobile readability", "No 320 px page overflow", "Fig. 6"),
    ]
    add_table(doc, ["Reported item", "Implemented response", "Evidence"], matrix, [2.0, 3.65, 0.85], font_size=8.8)

    doc.add_page_break()
    doc.add_heading("How the workflows operate", level=1)
    doc.add_heading("Department stock request", level=2)
    steps = [
        ("Create", "Marketing opens Warehouse > Fulfillment > Department requests and records purpose, cost center, required date, treatment, event, and items."),
        ("Approve", "The department approver confirms business need and budget; the requester sees their own status."),
        ("Fulfill", "Warehouse allocates, picks, scans, and releases approved stock."),
        ("Reconcile", "Event or department custody and Finance expense treatment remain attached to the request."),
    ]
    for i, (title, detail) in enumerate(steps, 1):
        add_step(doc, i, title, detail)
    doc.add_heading("Procurement receiving", level=2)
    steps = [
        ("Handoff", "Procurement issues an approved goods PO, which appears in the Operations receiving queue."),
        ("Compare", "Operations compares the PO, delivery evidence, and physical quantities."),
        ("Classify", "Each line is split into clean, damaged, unidentified, short, and excess quantities."),
        ("Control", "Clean stock proceeds; exception quantities generate controlled quality and supervisor work."),
    ]
    for i, (title, detail) in enumerate(steps, 1):
        add_step(doc, i, title, detail)
    doc.add_heading("Pick, pack, and release", level=2)
    steps = [
        ("Demand", "Orders enter from Ecommerce, Shopify, events, department requests, or controlled manual entry."),
        ("Allocate and scan", "Stock, serials, bundle members, rack/bin, and item scans are confirmed."),
        ("Pack", "Packaging supply, courier, waybill, and recipient/handover details are captured."),
        ("Close", "Release, failed delivery, proof of delivery, return, and Finance settlement remain traceable."),
    ]
    for i, (title, detail) in enumerate(steps, 1):
        add_step(doc, i, title, detail)

    doc.add_page_break()
    doc.add_heading("Live application evidence", level=1)
    add_body(doc, "Each figure below was captured from the deployed Vercel UAT application through a role-authenticated Playwright journey on August 26, 2026.")
    add_figure(doc, "01-live-purchase-order-receiving-queue.png", "Figure 1. Operations receiving queue", "PO 0001 and PO 0002 remain open with supplier, expected quantity, value, status, and Receive and inspect action. The page also explains the governed path from rider arrival to putaway.")
    add_figure(doc, "02-live-mixed-receipt-outcomes.png", "Figure 2. PO 0001 mixed receipt controls", "The Prodigy Watch line is staged with the exact stakeholder scenario: 50 clean, 20 damaged, 10 unidentified, 20 short, and 0 excess. The screenshot is evidence-only; the transaction was not submitted.")
    add_figure(doc, "03-live-marketing-stock-request.png", "Figure 3. Marketing’s own Event A request", "The Marketing Events Lead sees New stock request and the seeded Event A request with Marketing cost center, required date, expense treatment, and pending approval status.")
    add_figure(doc, "04-live-event-a-custody-and-reconciliation.png", "Figure 4. Event A lifecycle record", "Event A is present in the Events workspace beside other controlled records, ready for planning, warehouse handoff, custody, and reconciliation testing.")
    add_figure(doc, "05-live-pick-and-pack-scenario-queue.png", "Figure 5. Pick & Pack state-specific actions", "The same queue exposes event demand, a picking order with Confirm scanned pick, and a packing order with their current statuses and handoff data.")
    doc.add_page_break()
    add_figure(doc, "06-live-mobile-my-work-no-overflow.png", "Figure 6. 320 px Operations task view", "The long receipt identifier wraps within the task card. Bottom navigation remains reachable and neither the document nor body exceeds the viewport width.", width=3.05)

    doc.add_page_break()
    doc.add_heading("Validation record", level=1)
    validation_rows = [
        ("Scenario fixture tests", "7 / 7 passed", "Includes environment guard and Marketing ownership"),
        ("Live evidence journeys", "4 / 4 passed", "Operations, Marketing, Events, Pick & Pack, mobile"),
        ("Evidence reframing rerun", "3 / 3 passed", "Exact mixed values and scenario cards visible"),
        ("Mobile geometry", "Passed", "Document and body widths <= 320 px viewport"),
        ("Supabase readback", "Passed", "Counts and corrected Marketing owner confirmed"),
        ("Screenshot visual review", "6 / 6 accepted", "Legible, relevant, no obvious overlap or clipping"),
    ]
    add_table(doc, ["Control", "Result", "What it establishes"], validation_rows, [1.7, 1.2, 3.6], font_size=9.2)

    doc.add_heading("Recommended tester sequence", level=1)
    guidance = [
        "Marketing: open the seeded Event A request, create a second request, and confirm only authorized requests are visible.",
        "Operations: receive PO 0001 using the mixed outcome quantities, then verify quality exceptions and clean-stock putaway.",
        "Operations: receive PO 0002 as fulfillment supply and place it in F-01-02.",
        "Operations: exercise UAT-AUG24-PICKING and UAT-AUG24-PACKING, including scans, packaging, courier/waybill, and release.",
        "Operations Lead: review exceptions, split backorder behavior, and inventory/quality decisions.",
        "Finance and Marketing: reconcile Event A sales, giveaways, returns, losses, and expense treatment.",
        "Customer Service/Operations: use the four seeded return cases to validate serial lookup, replacement/refund, evidence, and closure.",
    ]
    for item in guidance:
        add_bullet(doc, item)

    add_status_callout(doc, "Important", "Transaction tests intentionally advance record state. Coordinate record ownership between testers. Rerun the guarded UAT seeder when deterministic starting records need to be restored; do not use it against production.", "caution")

    doc.add_heading("Traceability and limitations", level=1)
    add_body(doc, "Source report: wms comments (3).pdf, comments dated August 24 and August 25. Detailed engineering traceability remains in docs/audits/2026-08-25-WMS-AUG24-FEEDBACK-REVIEW.md. Evidence images are stored under docs/evidence/2026-08-26-aug24-25-remediation/.")
    add_body(doc, "This response certifies the reported August 24–25 items and supporting UAT fixtures. It is not a blanket production-readiness certification for future or unimplemented modules. New stakeholder feedback should include date, role, exact record, action, expected result, and observed result so it can be reproduced and added to regression coverage.")

    doc.core_properties.title = "mWell Intra WMS August 24–25 Remediation Response"
    doc.core_properties.subject = "UAT remediation evidence and tester readiness"
    doc.core_properties.author = "mWell Intra Project Team"
    doc.core_properties.keywords = "mWell Intra, WMS, UAT, remediation, evidence"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
